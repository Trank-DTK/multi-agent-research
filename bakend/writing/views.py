# 论文写作助手视图
import json
from django.http import JsonResponse, HttpResponse
from django.utils.encoding import smart_str
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser
from .models import Paper, PaperSection, Citation, WritingHistory
from agents.writing_agent import WritingService
from docx import Document
from io import BytesIO

writing_service = WritingService()


class GenerateOutlineView(APIView):
    """生成论文大纲接口"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        topic = request.data.get('topic')
        paper_type = request.data.get('paper_type', 'research')
        
        if not topic:
            return JsonResponse({'error': '主题不能为空'}, status=400)
        
        try:
            outline = writing_service.generate_outline(topic, paper_type)
            return JsonResponse({'outline': outline})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


class PaperCreateView(APIView):
    """创建论文接口"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        title = request.data.get('title')
        topic = request.data.get('topic')
        
        if not title:
            return JsonResponse({'error': '标题不能为空'}, status=400)
        
        # 创建论文
        paper = Paper.objects.create(
            user=request.user,
            title=title,
            status='draft'
        )
        
        # 生成大纲
        if topic:
            outline = writing_service.generate_outline(topic)
            # 解析大纲创建章节
            paper.outline = outline
            paper.save()
        
        return JsonResponse({
            'id': paper.id,
            'title': paper.title,
            'status': paper.status,
            'created_at': paper.created_at
        })


class PaperDetailView(APIView):
    """论文详情接口"""
    permission_classes = [IsAuthenticated]
    
    def get(self, request, paper_id):
        try:
            paper = Paper.objects.get(id=paper_id, user=request.user)
            sections = paper.sections.all()
            
            return JsonResponse({
                'id': paper.id,
                'title': paper.title,
                'abstract': paper.abstract,
                'keywords': paper.keywords,
                'content': paper.content,
                'status': paper.status,
                'sections': [{
                    'id': s.id,
                    'title': s.title,
                    'section_type': s.section_type,
                    'content': s.content,
                    'order': s.order
                } for s in sections],
                'created_at': paper.created_at,
                'updated_at': paper.updated_at
            })
        except Paper.DoesNotExist:
            return JsonResponse({'error': '论文不存在'}, status=404)
    
    def put(self, request, paper_id):
        try:
            paper = Paper.objects.get(id=paper_id, user=request.user)
            
            # 保存历史版本
            WritingHistory.objects.create(
                paper=paper,
                content=paper.content,
                version=WritingHistory.objects.filter(paper=paper).count() + 1
            )
            
            paper.title = request.data.get('title', paper.title)
            paper.abstract = request.data.get('abstract', paper.abstract)
            paper.keywords = request.data.get('keywords', paper.keywords)
            paper.content = request.data.get('content', paper.content)
            paper.status = request.data.get('status', paper.status)
            paper.save()
            
            return JsonResponse({'message': '保存成功'})
        except Paper.DoesNotExist:
            return JsonResponse({'error': '论文不存在'}, status=404)


class PaperListView(APIView):
    """论文列表接口"""
    permission_classes = [IsAuthenticated]
    
    def get(self, request):
        papers = Paper.objects.filter(user=request.user)
        return JsonResponse([{
            'id': p.id,
            'title': p.title,
            'status': p.status,
            'updated_at': p.updated_at
        } for p in papers], safe=False)


class PaperDeleteView(APIView):
    """删除论文接口"""
    permission_classes = [IsAuthenticated]
    
    def delete(self, request, paper_id):
        try:
            paper = Paper.objects.get(id=paper_id, user=request.user)
            paper.delete()
            return JsonResponse({'message': '删除成功'})
        except Paper.DoesNotExist:
            return JsonResponse({'error': '论文不存在'}, status=404)


class WriteSectionView(APIView):
    """撰写章节接口"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request, paper_id):
        section_title = request.data.get('title')
        context = request.data.get('context', '')
        word_count = request.data.get('word_count', 500)
        
        if not section_title:
            return JsonResponse({'error': '章节标题不能为空'}, status=400)
        
        try:
            content = writing_service.write_section(section_title, context, word_count)
            
            # 保存章节
            paper = Paper.objects.get(id=paper_id, user=request.user)
            section = PaperSection.objects.create(
                paper=paper,
                section_type='custom',
                title=section_title,
                content=content,
                order=paper.sections.count()
            )
            
            return JsonResponse({
                'section_id': section.id,
                'title': section.title,
                'content': content
            })
        except Paper.DoesNotExist:
            return JsonResponse({'error': '论文不存在'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


class PolishTextView(APIView):
    """润色文本接口"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        text = request.data.get('text')
        style = request.data.get('style', 'academic')
        
        if not text:
            return JsonResponse({'error': '文本不能为空'}, status=400)
        
        try:
            polished = writing_service.polish_text(text, style)
            return JsonResponse({'polished': polished})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


class GenerateAbstractView(APIView):
    """生成摘要接口"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request, paper_id):
        try:
            paper = Paper.objects.get(id=paper_id, user=request.user)
            content = paper.content or request.data.get('content', '')
            
            if not content:
                return JsonResponse({'error': '论文内容为空'}, status=400)
            
            abstract = writing_service.generate_abstract(content)
            
            # 保存摘要
            paper.abstract = abstract
            paper.save()
            
            return JsonResponse({'abstract': abstract})
        except Paper.DoesNotExist:
            return JsonResponse({'error': '论文不存在'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


class ExportDocxView(APIView):
    """导出Word文档接口"""
    permission_classes = [IsAuthenticated]
    
    def get(self, request, paper_id):
        try:
            paper = Paper.objects.get(id=paper_id, user=request.user)
            
            # 创建Word文档
            doc = Document()
            doc.add_heading(paper.title, 0)
            
            if paper.abstract:
                doc.add_heading('摘要', level=1)
                doc.add_paragraph(paper.abstract)
            
            # 添加章节
            sections = paper.sections.all().order_by('order')
            for section in sections:
                doc.add_heading(section.title, level=1)
                doc.add_paragraph(section.content)
            
            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # 返回文件
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            # 正确处理中文文件名
            filename = smart_str(paper.title)
            # 移除文件名中的非法字符
            import re
            filename = re.sub(r'[<>:"/\\|?*]', '', filename)
            # 添加文件扩展名
            filename = f'{filename}.docx'
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        except Paper.DoesNotExist:
            return JsonResponse({'error': '论文不存在'}, status=404)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'导出失败: {str(e)}'}, status=500)


class WritingAgentChatView(APIView):
    """写作Agent聊天接口"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        message = request.data.get('message')
        
        if not message:
            return JsonResponse({'error': '消息不能为空'}, status=400)
        
        try:
            response = writing_service.agent.run(message)
            return JsonResponse({'response': response})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
